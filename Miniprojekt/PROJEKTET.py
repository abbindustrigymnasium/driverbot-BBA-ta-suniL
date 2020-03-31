import openpyxl

wb = openpyxl.load_workbook('Bedömning_prog.xlsx')
sheet = wb.active

C_innehåll = wb.get_sheet_by_name('Centralt innehåll')
Kunskapskrav = wb.get_sheet_by_name('Kunskapskrav')
Kraven = wb.get_sheet_by_name('kraven')
Info = wb.get_sheet_by_name('Info')


sheet = wb["Centralt innehåll"]
sheet2 = wb["Kunskapskrav"]
sheet3 = wb["kraven"]
sheet4 = wb["Info"]

elever = [sheet2.cell(row=i, column=1).value for i in range(2, sheet2.max_row+1)]

kursinfo = {sheet4.cell(row=i, column=1).value:sheet4.cell(row=i, column=2).value for i in range(1, sheet4.max_row)}

resultat = [
    {sheet2.cell(row=e, column=1).value:
        {
            sheet2.cell(row=1, column=i).value:sheet2.cell(row=e, column=i).value for i in range(2, sheet2.max_column)
        }
    for e in range(2, sheet2.max_row)}
]

innehåll = [sheet.cell(row=1, column=i).value for i in range(1, sheet.max_column-1)] 

krav = [
    {sheet3.cell(row=e, column=1).value:
        {
            sheet3.cell(row=1, column=i).value:sheet3.cell(row=e, column=i).value for i in range(2, sheet3.max_column+1)
        }
    for e in range(2, sheet3.max_row)}
]

print('\n')
#print(kursinfo)
#print(elever)
print('\n')
#print(innehåll)
print(krav)


from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')